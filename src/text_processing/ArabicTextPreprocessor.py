import re
import json
from pathlib import Path
from typing import List, Dict, Any, Optional
import unicodedata
from dataclasses import dataclass
from datetime import datetime
import hashlib
import mammoth
from docx import Document

@dataclass
class TextChunk:
    """Data class to represent a text chunk with metadata"""
    id: str
    text: str
    cleaned_text: str
    metadata: Dict[str, Any]
    word_count: int
    char_count: int


class ArabicTextPreprocessor:
    """
    Comprehensive Arabic text preprocessing class for RAG applications.
    Handles text extraction, cleaning, chunking, and metadata generation.

    **NEW:**
    • Added removal of phone numbers, e‑mail addresses, and “flixat” (URLs/links).
    • Added three helper regex patterns and corresponding removal functions.
    • `clean_text()` now supports the optional flags:
        - `remove_phone_numbers`
        - `remove_emails`
        - `remove_flixat`
    """

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        """
        Initialize the preprocessor

        Args:
            chunk_size: Maximum number of words per chunk
            chunk_overlap: Number of words to overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        # Arabic diacritics (tashkeel) pattern
        self.diacritics_pattern = re.compile(r'[\u064B-\u0652\u0670\u0640]')

        # Tatweel (kashida) pattern
        self.tatweel_pattern = re.compile(r'\u0640+')

        # Multiple spaces/newlines pattern
        self.whitespace_pattern = re.compile(r'\s+')

        # Non‑Arabic characters (keeping numbers and basic punctuation)
        self.non_arabic_pattern = re.compile(
            r'[^\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF\s\d\.\،\؟\!\:\;\(\)\[\]\{\}\"\']+'
        )

        # Punctuation normalization
        self.punctuation_map = {
            '؟': '?',
            '،': ',',
            '؛': ';',
            '٪': '%',
            '٫': ',',
            '٬': ',',
            '۔': '.',
        }

        # === NEW REGEX PATTERNS ===
        # Phone numbers (handles international prefix, spaces, dashes, dots)
        self.phone_pattern = re.compile(
            r'\b(?:\+?[\d]{1,3}[-.\s]?)?(?:\(?\d{2,4}\)?[-.\s]?){1,4}\d{3,4}\b'
        )

        # E‑mail addresses
        self.email_pattern = re.compile(
            r'\b[\w\.-]+?@[\w\.-]+\.\w{2,}\b'
        )

        # “Flixat” → interpreted here as URLs / web links
        self.flixat_pattern = re.compile(
            r'(?:https?://\S+|www\.\S+)'
        )

    # ------------------------------------------------------------------
    #                        EXTRACTION METHODS
    # ------------------------------------------------------------------

    def extract_text_from_docx(self, file_path: str, use_uploaded_file: bool = True) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from DOCX file using multiple methods for better coverage

        Args:
            file_path: Path to the DOCX file (or filename if uploaded)
            use_uploaded_file: Whether to read from uploaded files via window.fs

        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            if use_uploaded_file:
                # Read uploaded file using window.fs.readFile
                try:
                    import js
                    file_data = js.window.fs.readFile(file_path)

                    # Convert to bytes for mammoth
                    import io
                    file_bytes = io.BytesIO(file_data.to_py())

                    # Method 1: Using mammoth (better formatting preservation)
                    result = mammoth.extract_raw_text(file_bytes)
                    text_mammoth = result.value

                    # Method 2: Using python-docx (backup method)
                    file_bytes.seek(0)  # Reset to beginning
                    doc = Document(file_bytes)
                    text_docx = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

                    extracted_text = text_mammoth if len(text_mammoth) > len(text_docx) else text_docx
                    file_size = len(file_data.to_py())

                except Exception as e:
                    print(f"Error with window.fs method, trying alternative approach: {e}")
                    # Fallback: try to read as if it's a local file
                    return self._extract_from_local_file(file_path)
            else:
                return self._extract_from_local_file(file_path)

            # Extract metadata
            metadata = {
                'source_file': file_path,
                'file_path': file_path,
                'extraction_method': 'mammoth' if len(text_mammoth) > len(text_docx) else 'python-docx',
                'extraction_timestamp': datetime.now().isoformat(),
                'original_length': len(extracted_text),
                'file_size_bytes': file_size,
            }

            # Try to extract document properties if available
            try:
                core_props = doc.core_properties
                if core_props.title:
                    metadata['title'] = core_props.title
                if core_props.author:
                    metadata['author'] = core_props.author
                if core_props.created:
                    metadata['created'] = core_props.created.isoformat()
                if core_props.modified:
                    metadata['modified'] = core_props.modified.isoformat()
            except:
                pass

            return extracted_text, metadata

        except Exception as e:
            raise Exception(f"Error extracting text from {file_path}: {str(e)}")

    def _extract_from_local_file(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract text from local file system"""
        # Method 1: Using mammoth (better formatting preservation)
        with open(file_path, "rb") as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            text_mammoth = result.value

        # Method 2: Using python-docx (backup method)
        doc = Document(file_path)
        text_docx = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

        # Use the longer text (usually mammoth gives better results)
        extracted_text = text_mammoth if len(text_mammoth) > len(text_docx) else text_docx

        # Extract metadata
        metadata = {
            'source_file': Path(file_path).name,
            'file_path': file_path,
            'extraction_method': 'mammoth' if len(text_mammoth) > len(text_docx) else 'python-docx',
            'extraction_timestamp': datetime.now().isoformat(),
            'original_length': len(extracted_text),
            'file_size_bytes': Path(file_path).stat().st_size,
        }

        # Try to extract document properties if available
        try:
            core_props = doc.core_properties
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.created:
                metadata['created'] = core_props.created.isoformat()
            if core_props.modified:
                metadata['modified'] = core_props.modified.isoformat()
        except:
            pass

        return extracted_text, metadata

    # ------------------------------------------------------------------
    #                        CLEANING HELPERS (NEW)
    # ------------------------------------------------------------------

    def remove_phone_numbers(self, text: str) -> str:
        """Remove phone numbers from text"""
        return self.phone_pattern.sub(' ', text)

    def remove_emails(self, text: str) -> str:
        """Remove e‑mail addresses from text"""
        return self.email_pattern.sub(' ', text)

    def remove_flixat(self, text: str) -> str:
        """Remove URLs / links (“flixat”) from text"""
        return self.flixat_pattern.sub(' ', text)

    # ------------------------------------------------------------------
    #                        EXISTING CLEANERS
    # ------------------------------------------------------------------

    def remove_diacritics(self, text: str) -> str:
        """Remove Arabic diacritics (tashkeel) from text"""
        return self.diacritics_pattern.sub('', text)

    def remove_tatweel(self, text: str) -> str:
        """Remove tatweel (kashida) characters"""
        return self.tatweel_pattern.sub('', text)

    def normalize_punctuation(self, text: str) -> str:
        """Normalize Arabic punctuation to standard forms"""
        for arabic_punct, standard_punct in self.punctuation_map.items():
            text = text.replace(arabic_punct, standard_punct)
        return text

    def remove_non_arabic(self, text: str) -> str:
        """Remove non-Arabic characters while preserving numbers and basic punctuation"""
        return self.non_arabic_pattern.sub(' ', text)

    def normalize_whitespace(self, text: str) -> str:
        """Normalize whitespace (multiple spaces, tabs, newlines)"""
        return self.whitespace_pattern.sub(' ', text).strip()

    # ------------------------------------------------------------------
    #                            CLEAN TEXT
    # ------------------------------------------------------------------

    def clean_text(
        self,
        text: str,
        remove_diacritics: bool = True,
        remove_tatweel: bool = True,
        normalize_punctuation: bool = True,
        remove_non_arabic: bool = False,
        normalize_whitespace: bool = True,
        # NEW flags
        remove_phone_numbers: bool = False,
        remove_emails: bool = False,
        remove_flixat: bool = False,
    ) -> str:
        """
        Apply comprehensive text cleaning

        Args:
            text: Input text
            remove_diacritics: Remove Arabic diacritics
            remove_tatweel: Remove tatweel characters
            normalize_punctuation: Normalize punctuation
            remove_non_arabic: Remove non-Arabic characters
            normalize_whitespace: Normalize whitespace
            remove_phone_numbers: Remove phone numbers
            remove_emails: Remove e‑mail addresses
            remove_flixat: Remove URLs / links

        Returns:
            Cleaned text
        """
        cleaned = text

        if remove_diacritics:
            cleaned = self.remove_diacritics(cleaned)

        if remove_tatweel:
            cleaned = self.remove_tatweel(cleaned)

        if normalize_punctuation:
            cleaned = self.normalize_punctuation(cleaned)

        if remove_non_arabic:
            cleaned = self.remove_non_arabic(cleaned)

        # --- NEW removals ---
        if remove_phone_numbers:
            cleaned = self.remove_phone_numbers(cleaned)

        if remove_emails:
            cleaned = self.remove_emails(cleaned)

        if remove_flixat:
            cleaned = self.remove_flixat(cleaned)

        if normalize_whitespace:
            cleaned = self.normalize_whitespace(cleaned)

        return cleaned

    # ------------------------------------------------------------------
    #        SENTENCE / PARAGRAPH SPLITTING + CHUNKING (unchanged)
    # ------------------------------------------------------------------

    def split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences using Arabic sentence delimiters"""
        sentence_endings = r'[.!?؟۔]+'
        sentences = re.split(sentence_endings, text)
        sentences = [s.strip() for s in sentences if s.strip()]
        return sentences

    def chunk_text(self, text: str, method: str = 'word_based') -> List[str]:
        """
        Split text into chunks for RAG processing

        Args:
            text: Input text
            method: Chunking method ('word_based', 'sentence_based', or 'paragraph_based')

        Returns:
            List of text chunks
        """
        if method == 'word_based':
            return self._word_based_chunking(text)
        elif method == 'sentence_based':
            return self._sentence_based_chunking(text)
        elif method == 'paragraph_based':
            return self._paragraph_based_chunking(text)
        else:
            raise ValueError(f"Unknown chunking method: {method}")

    def _word_based_chunking(self, text: str) -> List[str]:
        """Split text into word-based chunks with overlap"""
        words = text.split()
        chunks = []

        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = ' '.join(chunk_words)
            if chunk_text.strip():
                chunks.append(chunk_text)

        return chunks

    def _sentence_based_chunking(self, text: str) -> List[str]:
        """Split text into sentence-based chunks"""
        sentences = self.split_into_sentences(text)
        chunks = []
        current_chunk = []
        current_word_count = 0

        for sentence in sentences:
            sentence_words = len(sentence.split())

            if current_word_count + sentence_words > self.chunk_size and current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = current_chunk[-1:] if current_chunk else []
                current_word_count = len(' '.join(current_chunk).split())

            current_chunk.append(sentence)
            current_word_count += sentence_words

        if current_chunk:
            chunks.append(' '.join(current_chunk))

        return chunks

    def _paragraph_based_chunking(self, text: str) -> List[str]:
        """Split text into paragraph-based chunks"""
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        chunks = []
        current_chunk = []
        current_word_count = 0

        for paragraph in paragraphs:
            paragraph_words = len(paragraph.split())

            if current_word_count + paragraph_words > self.chunk_size and current_chunk:
                chunks.append('\n\n'.join(current_chunk))
                current_chunk = []
                current_word_count = 0

            current_chunk.append(paragraph)
            current_word_count += paragraph_words

        if current_chunk:
            chunks.append('\n\n'.join(current_chunk))

        return chunks

    # ------------------------------------------------------------------
    #                          UTILITIES
    # ------------------------------------------------------------------

    def generate_chunk_id(self, text: str, index: int) -> str:
        """Generate unique ID for chunk"""
        text_hash = hashlib.md5(text.encode('utf-8')).hexdigest()[:8]
        return f"chunk_{index:04d}_{text_hash}"

    def extract_keywords(self, text: str, top_k: int = 10) -> List[str]:
        """
        Extract potential keywords from Arabic text
        (Simple frequency-based approach - can be enhanced with NLP libraries)
        """
        cleaned = self.clean_text(text)
        words = cleaned.split()

        arabic_stopwords = {
            'في', 'من', 'إلى', 'على', 'عن', 'مع', 'هذا', 'هذه', 'ذلك', 'تلك',
            'التي', 'الذي', 'التي', 'الذين', 'اللتان', 'اللذان', 'اللواتي', 'اللاتي',
            'كان', 'كانت', 'يكون', 'تكون', 'أن', 'إن', 'لكن', 'ولكن', 'أم', 'أما',
            'لا', 'ما', 'لم', 'لن', 'قد', 'فقد', 'وقد', 'ند', 'عند', 'بعد', 'قبل'
        }

        word_freq: Dict[str, int] = {}
        for word in words:
            if len(word) > 2 and word not in arabic_stopwords:
                word_freq[word] = word_freq.get(word, 0) + 1

        keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        return [word for word, freq in keywords[:top_k]]

    # ------------------------------------------------------------------
    #                       MAIN PROCESSING PIPELINE
    # ------------------------------------------------------------------

    def process_document(
        self,
        file_path: str,
        chunking_method: str = 'sentence_based',
        cleaning_options: Optional[Dict[str, bool]] = None,
        use_uploaded_file: bool = True
    ) -> List[TextChunk]:
        """
        Complete document processing pipeline

        Args:
            file_path: Path to the document (or filename if uploaded)
            chunking_method: Method for text chunking
            cleaning_options: Dictionary of cleaning options
            use_uploaded_file: Whether to read from uploaded files

        Returns:
            List of TextChunk objects
        """
        # Default cleaning options
        if cleaning_options is None:
            cleaning_options = {
                'remove_diacritics': True,
                'remove_tatweel': True,
                'normalize_punctuation': True,
                'remove_non_arabic': False,
                'normalize_whitespace': True,
                # NEW flags (disabled by default to preserve behaviour)
                'remove_phone_numbers': False,
                'remove_emails': False,
                'remove_flixat': False
            }

        print(f"Processing document: {file_path}")

        # Step 1: Extract text
        print("1. Extracting text...")
        original_text, base_metadata = self.extract_text_from_docx(file_path, use_uploaded_file)
        print(f"   Extracted {len(original_text)} characters")

        # Step 2: Clean text
        print("2. Cleaning text...")
        cleaned_text = self.clean_text(original_text, **cleaning_options)
        print(f"   Cleaned text: {len(cleaned_text)} characters")

        # Step 3: Chunk text
        print(f"3. Chunking text using {chunking_method} method...")
        chunks = self.chunk_text(cleaned_text, method=chunking_method)
        print(f"   Created {len(chunks)} chunks")

        # Step 4: Create TextChunk objects with metadata
        print("4. Generating metadata...")
        text_chunks: List[TextChunk] = []

        for i, chunk_text in enumerate(chunks):
            chunk_id = self.generate_chunk_id(chunk_text, i)

            cleaned_chunk = self.clean_text(chunk_text, **cleaning_options)

            word_count = len(chunk_text.split())
            char_count = len(chunk_text)

            keywords = self.extract_keywords(cleaned_chunk, top_k=5)

            metadata = {
                **base_metadata,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunking_method': chunking_method,
                'cleaning_options': cleaning_options,
                'keywords': keywords,
                'position_start': sum(len(c.split()) for c in chunks[:i]),
                'position_end': sum(len(c.split()) for c in chunks[:i + 1]),
            }

            text_chunks.append(
                TextChunk(
                    id=chunk_id,
                    text=chunk_text,
                    cleaned_text=cleaned_chunk,
                    metadata=metadata,
                    word_count=word_count,
                    char_count=char_count
                )
            )

        print(f"✓ Processing complete: {len(text_chunks)} chunks created")
        return text_chunks

    # ------------------------------------------------------------------
    #                     SAVE / LOAD CHUNKS (unchanged)
    # ------------------------------------------------------------------

    def save_processed_chunks(self, chunks: List[TextChunk], output_path: str):
        """Save processed chunks to JSON file"""
        chunks_data = []
        for chunk in chunks:
            chunks_data.append({
                'id': chunk.id,
                'text': chunk.text,
                'cleaned_text': chunk.cleaned_text,
                'metadata': chunk.metadata,
                'word_count': chunk.word_count,
                'char_count': chunk.char_count
            })

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(chunks_data, f, ensure_ascii=False, indent=2)

        print(f"Chunks saved to: {output_path}")

    def load_processed_chunks(self, input_path: str) -> List[TextChunk]:
        """Load processed chunks from JSON file"""
        with open(input_path, 'r', encoding='utf-8') as f:
            chunks_data = json.load(f)

        chunks = []
        for data in chunks_data:
            chunks.append(
                TextChunk(
                    id=data['id'],
                    text=data['text'],
                    cleaned_text=data['cleaned_text'],
                    metadata=data['metadata'],
                    word_count=data['word_count'],
                    char_count=data['char_count']
                )
            )
        return chunks


# ----------------------------------------------------------------------
########################################################################################################################################## EXAMPLE USAGE (unchanged — optional demonstration)
# ----------------------------------------------------------------------
# if __name__ == "__main__":
#     preprocessor = ArabicTextPreprocessor(chunk_size=400, chunk_overlap=50)

#     file_path = "/content/فتح-المغيث-بشرح-ألفية-الحديث2.docx"

#     try:
#         chunks = preprocessor.process_document(
#             file_path=file_path,
#             chunking_method='sentence_based',
#             cleaning_options={
#                 'remove_diacritics': True,
#                 'remove_tatweel': True,
#                 'normalize_punctuation': True,
#                 'remove_non_arabic': False,
#                 'normalize_whitespace': True,
#                 # Enable new flags here if desired:
#                 'remove_phone_numbers': True,
#                 'remove_emails': True,
#                 'remove_flixat': True
#             },
#             use_uploaded_file=True
#         )

#         print(f"\n📊 Processing Results:")
#         print(f"   • Total chunks: {len(chunks)}")
#         print(f"   • Average words per chunk: {sum(c.word_count for c in chunks) / len(chunks):.1f}")
#         print(f"   • Total words: {sum(c.word_count for c in chunks)}")

#         if chunks:
#             sample = chunks[0]
#             print(f"\n📝 Sample Chunk (First):")
#             print(f"   • ID: {sample.id}")
#             print(f"   • Words: {sample.word_count}")
#             print(f"   • Keywords: {', '.join(sample.metadata.get('keywords', []))}")
#             print(f"   • Text preview: {sample.cleaned_text[:200]}...")
# ##############################################################################################################################
#         output_file = "processed_chunks 32.json"
#         preprocessor.save_processed_chunks(chunks, output_file)

#         print(f"\n✅ Processing complete! Chunks saved to '{output_file}'")
#         print("💡 You can now use these chunks in your RAG model")

#     except FileNotFoundError:
#         print(f"❌ Error: File '{file_path}' not found")
#         print("📌 Make sure the file is in the same directory as this script")
#     except Exception as e:
#         print(f"❌ Error processing document: {str(e)}")
